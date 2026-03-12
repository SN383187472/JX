"""
报告生成模块 - 专业版样式
支持生成精美的Word和Excel格式报告
"""
from datetime import datetime
from typing import Dict, Any, List, Optional
import json
import os
import io
import logging

logger = logging.getLogger("report")


class ReportGenerator:
    """
    专业报告生成器
    
    生成符合安全基线核查标准的专业报告
    """
    
    # 类型名称映射
    TYPE_NAMES = {
        'kylin': '麒麟操作系统',
        'dameng': '达梦数据库',
        'windows7': 'Windows 7',
        'windows10': 'Windows 10',
        'windows2012': 'Windows Server 2012',
        'centos': 'CentOS',
        'ubuntu': 'Ubuntu'
    }
    
    # 严重程度名称
    SEVERITY_NAMES = {
        'high': '高',
        'medium': '中',
        'low': '低'
    }
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or "reports"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_json_report(
        self,
        result_data: Dict[str, Any],
        filename: str = None
    ) -> str:
        """生成JSON格式报告"""
        if not filename:
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        filepath = os.path.join(self.output_dir, f"{filename}.json")
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(result_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"JSON报告已生成: {filepath}")
        return filepath
    
    def generate_excel_report(
        self,
        result_data: Dict[str, Any],
        filename: str = None
    ) -> bytes:
        """
        生成专业Excel格式报告
        """
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side, NamedStyle
            from openpyxl.utils import get_column_letter
            from openpyxl.chart import PieChart, Reference
        except ImportError:
            logger.error("openpyxl未安装")
            return b""
        
        if not filename:
            filename = f"baseline_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        wb = openpyxl.Workbook()
        
        # ============ 样式定义 ============
        # 颜色方案
        COLORS = {
            'primary': '1F4E79',      # 深蓝色-主色
            'secondary': '2E75B6',    # 蓝色-辅助
            'success': '70AD47',      # 绿色-通过
            'danger': 'C00000',       # 红色-不符合
            'warning': 'FFC000',      # 黄色-警告
            'light': 'F2F2F2',        # 浅灰色-背景
            'white': 'FFFFFF',
            'black': '000000'
        }
        
        # 字体样式
        title_font = Font(name='微软雅黑', size=20, bold=True, color=COLORS['primary'])
        subtitle_font = Font(name='微软雅黑', size=14, bold=True, color=COLORS['secondary'])
        header_font = Font(name='微软雅黑', size=11, bold=True, color=COLORS['white'])
        normal_font = Font(name='微软雅黑', size=10, color=COLORS['black'])
        bold_font = Font(name='微软雅黑', size=10, bold=True, color=COLORS['black'])
        
        # 填充样式
        header_fill = PatternFill(start_color=COLORS['primary'], end_color=COLORS['primary'], fill_type='solid')
        pass_fill = PatternFill(start_color=COLORS['success'], end_color=COLORS['success'], fill_type='solid')
        fail_fill = PatternFill(start_color=COLORS['danger'], end_color=COLORS['danger'], fill_type='solid')
        warning_fill = PatternFill(start_color=COLORS['warning'], end_color=COLORS['warning'], fill_type='solid')
        light_fill = PatternFill(start_color=COLORS['light'], end_color=COLORS['light'], fill_type='solid')
        
        # 边框样式
        thin_border = Border(
            left=Side(style='thin', color='D0D0D0'),
            right=Side(style='thin', color='D0D0D0'),
            top=Side(style='thin', color='D0D0D0'),
            bottom=Side(style='thin', color='D0D0D0')
        )
        
        # ============ Sheet 1: 概览 ============
        ws_summary = wb.active
        ws_summary.title = "概览"
        
        # 标题区域
        ws_summary.merge_cells('A1:H1')
        ws_summary['A1'] = "安全基线核查报告"
        ws_summary['A1'].font = title_font
        ws_summary['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws_summary.row_dimensions[1].height = 40
        
        ws_summary.merge_cells('A2:H2')
        ws_summary['A2'] = f"Security Baseline Assessment Report"
        ws_summary['A2'].font = Font(name='Arial', size=12, italic=True, color='808080')
        ws_summary['A2'].alignment = Alignment(horizontal='center')
        
        # 报告信息区域
        ws_summary.merge_cells('A4:B4')
        ws_summary['A4'] = "报告信息"
        ws_summary['A4'].font = subtitle_font
        
        info_labels = ['目标名称', '目标类型', '扫描时间', '规则总数', '通过规则', '不符合规则', '合规率']
        target_name = result_data.get('target_name', '未知')
        target_type = self.TYPE_NAMES.get(result_data.get('target_type', ''), result_data.get('target_type', '未知'))
        scan_time = result_data.get('scan_time', '')
        total = result_data.get('total_rules', 0)
        passed = result_data.get('passed_rules', 0)
        failed = result_data.get('failed_rules', 0)
        pass_rate = f"{passed / max(total, 1) * 100:.1f}%"
        
        info_values = [target_name, target_type, scan_time, str(total), str(passed), str(failed), pass_rate]
        
        for i, (label, value) in enumerate(zip(info_labels, info_values)):
            row = 5 + i
            ws_summary['A' + str(row)] = label
            ws_summary['A' + str(row)].font = bold_font
            ws_summary['A' + str(row)].fill = light_fill
            ws_summary['A' + str(row)].border = thin_border
            
            ws_summary.merge_cells(f'B{row}:D{row}')
            ws_summary['B' + str(row)] = value
            ws_summary['B' + str(row)].font = normal_font
            ws_summary['B' + str(row)].border = thin_border
        
        # 检查结果汇总
        ws_summary.merge_cells('A13:D13')
        ws_summary['A13'] = "检查结果汇总"
        ws_summary['A13'].font = subtitle_font
        
        # 统计表格
        headers = ['结果状态', '数量', '占比']
        for col, header in enumerate(headers, 1):
            cell = ws_summary.cell(row=14, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        stats = [
            ('符合', passed, f"{passed / max(total, 1) * 100:.1f}%"),
            ('不符合', failed, f"{failed / max(total, 1) * 100:.1f}%")
        ]
        
        for row_idx, (status, count, rate) in enumerate(stats, 15):
            ws_summary.cell(row=row_idx, column=1, value=status).border = thin_border
            ws_summary.cell(row=row_idx, column=2, value=count).border = thin_border
            ws_summary.cell(row=row_idx, column=3, value=rate).border = thin_border
            
            if status == '符合':
                ws_summary.cell(row=row_idx, column=1).fill = pass_fill
                ws_summary.cell(row=row_idx, column=1).font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
            else:
                ws_summary.cell(row=row_idx, column=1).fill = fail_fill
                ws_summary.cell(row=row_idx, column=1).font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
        
        # 按严重程度统计
        ws_summary.merge_cells('A18:D18')
        ws_summary['A18'] = "按严重程度统计"
        ws_summary['A18'].font = subtitle_font
        
        severity_headers = ['严重程度', '总数', '通过', '不符合']
        for col, header in enumerate(severity_headers, 1):
            cell = ws_summary.cell(row=19, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        severity_stats = {'high': {'pass': 0, 'fail': 0}, 'medium': {'pass': 0, 'fail': 0}, 'low': {'pass': 0, 'fail': 0}}
        for detail in result_data.get('details', []):
            sev = detail.get('severity', 'medium')
            if sev in severity_stats:
                if detail.get('status') == 'pass':
                    severity_stats[sev]['pass'] += 1
                else:
                    severity_stats[sev]['fail'] += 1
        
        for row_idx, sev in enumerate(['high', 'medium', 'low'], 20):
            sev_name = self.SEVERITY_NAMES.get(sev, sev)
            ws_summary.cell(row=row_idx, column=1, value=sev_name).border = thin_border
            total_sev = severity_stats[sev]['pass'] + severity_stats[sev]['fail']
            ws_summary.cell(row=row_idx, column=2, value=total_sev).border = thin_border
            ws_summary.cell(row=row_idx, column=3, value=severity_stats[sev]['pass']).border = thin_border
            ws_summary.cell(row=row_idx, column=4, value=severity_stats[sev]['fail']).border = thin_border
        
        # 按类别统计
        ws_summary.merge_cells('A24:D24')
        ws_summary['A24'] = "按类别统计"
        ws_summary['A24'].font = subtitle_font
        
        cat_headers = ['类别', '总数', '通过', '不符合']
        for col, header in enumerate(cat_headers, 1):
            cell = ws_summary.cell(row=25, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        
        categories = {}
        for detail in result_data.get('details', []):
            cat = detail.get('category', '其他')
            if cat not in categories:
                categories[cat] = {'pass': 0, 'fail': 0}
            if detail.get('status') == 'pass':
                categories[cat]['pass'] += 1
            else:
                categories[cat]['fail'] += 1
        
        for row_idx, (cat, stat) in enumerate(categories.items(), 26):
            ws_summary.cell(row=row_idx, column=1, value=cat).border = thin_border
            ws_summary.cell(row=row_idx, column=2, value=stat['pass'] + stat['fail']).border = thin_border
            ws_summary.cell(row=row_idx, column=3, value=stat['pass']).border = thin_border
            ws_summary.cell(row=row_idx, column=4, value=stat['fail']).border = thin_border
        
        # 设置列宽
        ws_summary.column_dimensions['A'].width = 15
        ws_summary.column_dimensions['B'].width = 20
        ws_summary.column_dimensions['C'].width = 12
        ws_summary.column_dimensions['D'].width = 12
        
        # ============ Sheet 2: 详细结果 ============
        ws_detail = wb.create_sheet("详细结果")
        
        # 标题
        ws_detail.merge_cells('A1:K1')
        ws_detail['A1'] = f"详细检查结果 - {target_name}"
        ws_detail['A1'].font = Font(name='微软雅黑', size=16, bold=True, color=COLORS['primary'])
        ws_detail['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws_detail.row_dimensions[1].height = 30
        
        # 表头
        headers = ['序号', '规则ID', '规则名称', '描述', '类别', '严重程度', '状态', '检查命令', '期望结果', '实际输出', '分析说明']
        for col, header in enumerate(headers, 1):
            cell = ws_detail.cell(row=3, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border
        
        # 数据行
        details = result_data.get('details', [])
        for row_idx, detail in enumerate(details, 4):
            status = detail.get('status', 'unknown')
            severity = detail.get('severity', 'medium')
            
            if status == 'pass':
                status_text = '符合'
                status_fill = pass_fill
            elif status == 'fail':
                status_text = '不符合'
                status_fill = fail_fill
            else:
                status_text = '错误'
                status_fill = warning_fill
            
            row_data = [
                row_idx - 3,
                detail.get('rule_id', ''),
                detail.get('rule_name', ''),
                detail.get('description', ''),
                detail.get('category', ''),
                self.SEVERITY_NAMES.get(severity, severity),
                status_text,
                detail.get('command', ''),
                detail.get('expected', ''),
                (detail.get('output', '') or '')[:200],
                detail.get('analysis', '')
            ]
            
            for col, value in enumerate(row_data, 1):
                cell = ws_detail.cell(row=row_idx, column=col, value=value)
                cell.font = normal_font
                cell.border = thin_border
                cell.alignment = Alignment(vertical='center', wrap_text=True)
                
                if col == 6:  # 严重程度
                    if severity == 'high':
                        cell.fill = PatternFill(start_color='FFCDD2', end_color='FFCDD2', fill_type='solid')
                    elif severity == 'medium':
                        cell.fill = PatternFill(start_color='FFF9C4', end_color='FFF9C4', fill_type='solid')
                
                if col == 7:  # 状态
                    cell.fill = status_fill
                    cell.font = Font(name='微软雅黑', size=10, bold=True, color='FFFFFF')
                    cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 设置列宽
        col_widths = [6, 12, 20, 25, 12, 10, 8, 30, 15, 30, 25]
        for i, width in enumerate(col_widths, 1):
            ws_detail.column_dimensions[get_column_letter(i)].width = width
        
        # 冻结表头
        ws_detail.freeze_panes = 'A4'
        
        # 保存
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        logger.info("Excel报告已生成")
        return output.getvalue()
    
    def generate_word_report(
        self,
        result_data: Dict[str, Any],
        filename: str = None
    ) -> bytes:
        """
        生成专业Word格式报告
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            logger.error("python-docx未安装")
            return b""
        
        if not filename:
            filename = f"baseline_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        doc = Document()
        
        # 设置默认字体
        doc.styles['Normal'].font.name = '微软雅黑'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # ============ 封面 ============
        # 空行
        for _ in range(3):
            doc.add_paragraph()
        
        # 标题
        title = doc.add_paragraph()
        title_run = title.add_run("安全基线核查报告")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 121)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run("Security Baseline Assessment Report")
        sub_run.font.size = Pt(14)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(128, 128, 128)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 空行
        for _ in range(5):
            doc.add_paragraph()
        
        # 基本信息
        target_name = result_data.get('target_name', '未知')
        target_type = self.TYPE_NAMES.get(result_data.get('target_type', ''), result_data.get('target_type', '未知'))
        scan_time = result_data.get('scan_time', '')
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"目标名称: {target_name}\n").font.size = Pt(14)
        info_para.add_run(f"目标类型: {target_type}\n").font.size = Pt(14)
        info_para.add_run(f"扫描时间: {scan_time}\n").font.size = Pt(14)
        
        # 分页
        doc.add_page_break()
        
        # ============ 目录 ============
        doc.add_heading('目录', level=1)
        toc_items = [
            '一、报告概述',
            '二、检查结果统计',
            '三、详细检查结果',
            '四、风险分析及建议'
        ]
        for item in toc_items:
            p = doc.add_paragraph(item)
            p.paragraph_format.left_indent = Cm(1)
        
        doc.add_page_break()
        
        # ============ 一、报告概述 ============
        doc.add_heading('一、报告概述', level=1)
        
        total = result_data.get('total_rules', 0)
        passed = result_data.get('passed_rules', 0)
        failed = result_data.get('failed_rules', 0)
        pass_rate = passed / max(total, 1) * 100
        
        # 概述段落
        overview = doc.add_paragraph()
        overview.add_run(f"本次安全基线核查针对 {target_name}({target_type}) 进行，共检查 {total} 项安全基线要求。")
        overview.add_run(f"\n\n检查结果：通过 {passed} 项，不符合 {failed} 项，合规率 {pass_rate:.1f}%。")
        
        if failed > 0:
            risk_para = doc.add_paragraph()
            risk_run = risk_para.add_run(f"\n⚠ 存在 {failed} 项不符合要求，建议尽快整改。")
            risk_run.font.color.rgb = RGBColor(192, 0, 0)
            risk_run.font.bold = True
        
        # ============ 二、检查结果统计 ============
        doc.add_heading('二、检查结果统计', level=1)
        
        # 2.1 总体情况
        doc.add_heading('2.1 总体情况', level=2)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('检查项目', str(total)),
            ('通过项目', str(passed)),
            ('不符合项目', str(failed)),
            ('合规率', f'{pass_rate:.1f}%'),
            ('检查结论', '符合要求' if pass_rate >= 90 else ('基本符合' if pass_rate >= 70 else '不符合要求'))
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
        
        # 2.2 按严重程度统计
        doc.add_heading('2.2 按严重程度统计', level=2)
        
        severity_stats = {'高': {'pass': 0, 'fail': 0}, '中': {'pass': 0, 'fail': 0}, '低': {'pass': 0, 'fail': 0}}
        severity_map = {'high': '高', 'medium': '中', 'low': '低'}
        
        for detail in result_data.get('details', []):
            sev = severity_map.get(detail.get('severity', 'medium'), '中')
            if detail.get('status') == 'pass':
                severity_stats[sev]['pass'] += 1
            else:
                severity_stats[sev]['fail'] += 1
        
        sev_table = doc.add_table(rows=4, cols=4)
        sev_table.style = 'Table Grid'
        
        sev_table.cell(0, 0).text = '严重程度'
        sev_table.cell(0, 1).text = '总数'
        sev_table.cell(0, 2).text = '通过'
        sev_table.cell(0, 3).text = '不符合'
        
        for i, sev in enumerate(['高', '中', '低'], 1):
            sev_table.cell(i, 0).text = sev
            total_sev = severity_stats[sev]['pass'] + severity_stats[sev]['fail']
            sev_table.cell(i, 1).text = str(total_sev)
            sev_table.cell(i, 2).text = str(severity_stats[sev]['pass'])
            sev_table.cell(i, 3).text = str(severity_stats[sev]['fail'])
        
        # 2.3 按类别统计
        doc.add_heading('2.3 按类别统计', level=2)
        
        categories = {}
        for detail in result_data.get('details', []):
            cat = detail.get('category', '其他')
            if cat not in categories:
                categories[cat] = {'pass': 0, 'fail': 0}
            if detail.get('status') == 'pass':
                categories[cat]['pass'] += 1
            else:
                categories[cat]['fail'] += 1
        
        cat_table = doc.add_table(rows=len(categories) + 1, cols=4)
        cat_table.style = 'Table Grid'
        
        cat_table.cell(0, 0).text = '类别'
        cat_table.cell(0, 1).text = '总数'
        cat_table.cell(0, 2).text = '通过'
        cat_table.cell(0, 3).text = '不符合'
        
        for i, (cat, stat) in enumerate(categories.items(), 1):
            cat_table.cell(i, 0).text = cat
            cat_table.cell(i, 1).text = str(stat['pass'] + stat['fail'])
            cat_table.cell(i, 2).text = str(stat['pass'])
            cat_table.cell(i, 3).text = str(stat['fail'])
        
        # ============ 三、详细检查结果 ============
        doc.add_heading('三、详细检查结果', level=1)
        
        # 不符合项
        failed_items = [d for d in result_data.get('details', []) if d.get('status') == 'fail']
        if failed_items:
            doc.add_heading('3.1 不符合项', level=2)
            
            for detail in failed_items:
                # 规则标题
                p = doc.add_paragraph()
                run = p.add_run(f"【{detail.get('rule_id', '')}】{detail.get('rule_name', '')}")
                run.font.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)
                
                # 详情表格
                detail_table = doc.add_table(rows=5, cols=2)
                detail_table.style = 'Table Grid'
                
                detail_data = [
                    ('描述', detail.get('description', '')),
                    ('检查命令', detail.get('command', '')),
                    ('期望结果', detail.get('expected', '')),
                    ('实际输出', (detail.get('output', '') or '')[:300]),
                    ('分析说明', detail.get('analysis', ''))
                ]
                
                for j, (label, value) in enumerate(detail_data):
                    detail_table.cell(j, 0).text = label
                    detail_table.cell(j, 1).text = value
                
                doc.add_paragraph()
        
        # 符合项
        passed_items = [d for d in result_data.get('details', []) if d.get('status') == 'pass']
        if passed_items:
            doc.add_heading('3.2 符合项', level=2)
            
            # 简洁表格
            pass_table = doc.add_table(rows=len(passed_items) + 1, cols=4)
            pass_table.style = 'Table Grid'
            
            pass_table.cell(0, 0).text = '规则ID'
            pass_table.cell(0, 1).text = '规则名称'
            pass_table.cell(0, 2).text = '类别'
            pass_table.cell(0, 3).text = '严重程度'
            
            for i, detail in enumerate(passed_items, 1):
                pass_table.cell(i, 0).text = detail.get('rule_id', '')
                pass_table.cell(i, 1).text = detail.get('rule_name', '')
                pass_table.cell(i, 2).text = detail.get('category', '')
                pass_table.cell(i, 3).text = self.SEVERITY_NAMES.get(detail.get('severity', 'medium'), '中')
        
        # ============ 四、风险分析及建议 ============
        doc.add_heading('四、风险分析及建议', level=1)
        
        if failed > 0:
            doc.add_paragraph(f"本次检查发现 {failed} 项不符合要求，建议采取以下措施：")
            doc.add_paragraph("1. 对高风险不符合项优先整改")
            doc.add_paragraph("2. 分析不符合项的根本原因")
            doc.add_paragraph("3. 制定整改计划并跟踪落实")
            doc.add_paragraph("4. 整改完成后进行复查验证")
        else:
            doc.add_paragraph("本次检查所有项目均符合要求，建议继续保持。")
        
        # 保存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        logger.info("Word报告已生成")
        return output.getvalue()
    
    def generate_summary(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """生成摘要统计"""
        total_rules = sum(r.get("total_rules", 0) for r in results)
        passed_rules = sum(r.get("passed_rules", 0) for r in results)
        failed_rules = sum(r.get("failed_rules", 0) for r in results)
        
        return {
            "total_scans": len(results),
            "total_rules": total_rules,
            "passed_rules": passed_rules,
            "failed_rules": failed_rules,
            "pass_rate": round(passed_rules / total_rules * 100, 2) if total_rules > 0 else 0,
            "generated_at": datetime.now().isoformat()
        }
